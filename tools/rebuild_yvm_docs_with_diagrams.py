# -*- coding: utf-8 -*-
from pathlib import Path
from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs"
ASSET_DIR = OUT / "assets"
OUT.mkdir(exist_ok=True)
ASSET_DIR.mkdir(exist_ok=True)

ACCENT = "1F4E79"
LIGHT = "EAF2F8"
MID = "D9EAF7"
GRAY = "F5F7F9"


def font(size, bold=False):
    candidates = [
        r"C:\Windows\Fonts\msyhbd.ttc" if bold else r"C:\Windows\Fonts\msyh.ttc",
        r"C:\Windows\Fonts\simhei.ttf",
        r"C:\Windows\Fonts\simsun.ttc",
    ]
    for path in candidates:
        try:
            return ImageFont.truetype(path, size)
        except Exception:
            pass
    return ImageFont.load_default()


def wrap_text(draw, text, fnt, width):
    lines, current = [], ""
    for ch in text:
        test = current + ch
        if draw.textbbox((0, 0), test, font=fnt)[2] <= width:
            current = test
        else:
            if current:
                lines.append(current)
            current = ch
    if current:
        lines.append(current)
    return lines


def rounded_box(draw, box, title, body="", fill="#EAF2F8", outline="#7EA6C7"):
    x1, y1, x2, y2 = box
    draw.rounded_rectangle(box, radius=16, fill=fill, outline=outline, width=2)
    title_font = font(28, True)
    body_font = font(21)
    draw.text((x1 + 18, y1 + 14), title, font=title_font, fill="#173B5B")
    if body:
        y = y1 + 54
        for line in wrap_text(draw, body, body_font, x2 - x1 - 36):
            draw.text((x1 + 18, y), line, font=body_font, fill="#2D3A45")
            y += 28


def arrow(draw, start, end, color="#315F86"):
    draw.line([start, end], fill=color, width=4)
    ex, ey = end
    sx, sy = start
    if abs(ex - sx) >= abs(ey - sy):
        sign = 1 if ex > sx else -1
        pts = [(ex, ey), (ex - sign * 18, ey - 9), (ex - sign * 18, ey + 9)]
    else:
        sign = 1 if ey > sy else -1
        pts = [(ex, ey), (ex - 9, ey - sign * 18), (ex + 9, ey - sign * 18)]
    draw.polygon(pts, fill=color)


def make_architecture(path):
    img = Image.new("RGB", (1500, 900), "white")
    d = ImageDraw.Draw(img)
    d.text((50, 34), "YVM-SLAM 总体架构", font=font(36, True), fill="#173B5B")
    boxes = {
        "data": (70, 130, 360, 280),
        "ros": (70, 350, 360, 520),
        "system": (480, 150, 760, 300),
        "tracking": (880, 130, 1210, 310),
        "orb": (880, 380, 1210, 550),
        "map": (480, 600, 760, 760),
        "backend": (880, 620, 1210, 780),
        "out": (1260, 260, 1450, 600),
    }
    rounded_box(d, boxes["data"], "离线数据", "TUM RGB-D\nRGB / Depth / timestamp", "#F7FBFF")
    rounded_box(d, boxes["ros"], "ROS2 语义输入", "/dynamic_mask\n/dynamic_detections", "#F2F8F2", "#7CB585")
    rounded_box(d, boxes["system"], "System", "加载词袋和配置\n创建线程与全局模块", "#EAF2F8")
    rounded_box(d, boxes["tracking"], "Tracking", "GrabImageRGBD\nTrack\n对象更新", "#EAF2F8")
    rounded_box(d, boxes["orb"], "ORBextractor", "FAST/ORB 特征\n按 mask/bbox 删除动态点", "#FFF8E8", "#D5A84A")
    rounded_box(d, boxes["map"], "Atlas / Map", "KeyFrame\nMapPoint\nMapObject", "#F4ECF7", "#A678B4")
    rounded_box(d, boxes["backend"], "后台线程", "LocalMapping\nLoopClosing\nObjectOptimizer\nViewer", "#F4ECF7", "#A678B4")
    rounded_box(d, boxes["out"], "输出", "轨迹文件\nViewer\n/odom 与 TF", "#F2F8F2", "#7CB585")
    arrow(d, (360, 205), (480, 220))
    arrow(d, (360, 435), (480, 245))
    arrow(d, (760, 225), (880, 220))
    arrow(d, (1045, 310), (1045, 380))
    arrow(d, (880, 465), (760, 680))
    arrow(d, (760, 680), (880, 700))
    arrow(d, (1210, 220), (1260, 330))
    arrow(d, (1210, 700), (1260, 500))
    img.save(path)


def make_data_flow(path):
    img = Image.new("RGB", (1500, 760), "white")
    d = ImageDraw.Draw(img)
    d.text((50, 32), "单帧 RGB-D 处理流程", font=font(36, True), fill="#173B5B")
    steps = [
        ("读入帧", "RGB、Depth、timestamp"),
        ("接收语义", "最新 mask / detections"),
        ("动态过滤", "YoloDetection -> ORBextractor"),
        ("构建 Frame", "灰度图、深度、ORB 特征"),
        ("SLAM 跟踪", "运动模型、局部地图、重定位"),
        ("对象更新", "静态 bbox + 深度 -> MapObject"),
        ("输出", "Tcw、轨迹、Viewer、/odom"),
    ]
    x, y = 45, 170
    w, h = 185, 180
    for i, (title, body) in enumerate(steps):
        rounded_box(d, (x + i * 205, y, x + i * 205 + w, y + h), title, body, "#EAF2F8" if i % 2 == 0 else "#F7FBFF")
        if i < len(steps) - 1:
            arrow(d, (x + i * 205 + w, y + h // 2), (x + (i + 1) * 205 - 10, y + h // 2))
    d.text((80, 470), "关键点：mask 不是在优化阶段才处理，而是在 ORB 特征提取阶段直接删除动态区域关键点。这样后续匹配、位姿估计和建图尽量只依赖静态背景。", font=font(25), fill="#2D3A45")
    img.save(path)


def make_run_workflow(path):
    img = Image.new("RGB", (1400, 880), "white")
    d = ImageDraw.Draw(img)
    d.text((50, 34), "推荐运行流程", font=font(36, True), fill="#173B5B")
    boxes = [
        ((90, 130, 430, 270), "1. 准备依赖", "OpenCV / Pangolin / GTSAM\nONNX Runtime / ROS2"),
        ((530, 130, 870, 270), "2. 构建项目", "./build.sh\n或 cmake + make rgbd_tum"),
        ((970, 130, 1310, 270), "3. 离线基线", "先不带 --ros\n确认轨迹能生成"),
        ((970, 420, 1310, 560), "4. 启用 ROS2", "发布图像、深度、/odom\n订阅 mask/detections"),
        ((530, 420, 870, 560), "5. 接入分割节点", "检查 /dynamic_mask\n确认尺寸和编码"),
        ((90, 420, 430, 560), "6. 评估结果", "ATE、跟踪时间\nViewer 截图与日志"),
        ((530, 650, 870, 800), "7. 调参数", "DepthMapFactor\nORB 参数、mask 阈值"),
    ]
    for b, title, body in boxes:
        rounded_box(d, b, title, body, "#F7FBFF")
    arrow(d, (430, 200), (530, 200))
    arrow(d, (870, 200), (970, 200))
    arrow(d, (1140, 270), (1140, 420))
    arrow(d, (970, 490), (870, 490))
    arrow(d, (530, 490), (430, 490))
    arrow(d, (260, 560), (530, 725))
    arrow(d, (870, 725), (1140, 560))
    img.save(path)


def make_troubleshooting(path):
    img = Image.new("RGB", (1400, 900), "white")
    d = ImageDraw.Draw(img)
    d.text((50, 34), "问题排查路径", font=font(36, True), fill="#173B5B")
    nodes = [
        ((520, 110, 880, 240), "程序跑不起来", "先看错误发生在哪一步"),
        ((90, 330, 410, 470), "构建失败", "检查依赖、ONNXRUNTIME_DIR、ROS2 source"),
        ((540, 330, 860, 470), "运行失败", "检查词袋、YAML、数据路径、association"),
        ((990, 330, 1310, 470), "效果不好", "检查内参、深度尺度、mask 对齐"),
        ((90, 610, 410, 760), "链接错误", "库路径和版本不一致"),
        ((540, 610, 860, 760), "图像读取错误", "路径拼接或文件缺失"),
        ((990, 610, 1310, 760), "动态过滤异常", "mask 尺寸、编码、时间同步"),
    ]
    for b, title, body in nodes:
        rounded_box(d, b, title, body, "#FFF8E8" if "失败" in title or "异常" in title else "#EAF2F8")
    arrow(d, (610, 240), (250, 330))
    arrow(d, (700, 240), (700, 330))
    arrow(d, (790, 240), (1150, 330))
    arrow(d, (250, 470), (250, 610))
    arrow(d, (700, 470), (700, 610))
    arrow(d, (1150, 470), (1150, 610))
    img.save(path)


def make_test_loop(path):
    img = Image.new("RGB", (1400, 820), "white")
    d = ImageDraw.Draw(img)
    d.text((50, 34), "测试闭环", font=font(36, True), fill="#173B5B")
    cx, cy = 700, 430
    boxes = [
        ((520, 100, 880, 230), "构建验证", "cmake / make / 依赖检查"),
        ((920, 300, 1280, 430), "离线基线", "TUM 序列 + 轨迹文件"),
        ((920, 560, 1280, 700), "语义集成", "mask / detections / 对象地图"),
        ((520, 610, 880, 750), "轨迹评估", "ATE、跟踪时间、LOST 次数"),
        ((120, 560, 480, 700), "缺陷记录", "复现命令、日志、截图"),
        ((120, 300, 480, 430), "回归复测", "固定数据集重复验证"),
    ]
    for b, title, body in boxes:
        rounded_box(d, b, title, body, "#F7FBFF")
    arrow(d, (700, 230), (920, 350))
    arrow(d, (1100, 430), (1100, 560))
    arrow(d, (920, 650), (880, 690))
    arrow(d, (520, 690), (480, 650))
    arrow(d, (300, 560), (300, 430))
    arrow(d, (480, 365), (520, 165))
    d.ellipse((610, 340, 790, 520), fill="#EAF2F8", outline="#7EA6C7", width=3)
    d.text((648, 385), "持续\n改进", font=font(30, True), fill="#173B5B")
    img.save(path)


def generate_diagrams():
    paths = {
        "architecture": ASSET_DIR / "architecture.png",
        "data_flow": ASSET_DIR / "data_flow.png",
        "run_workflow": ASSET_DIR / "run_workflow.png",
        "troubleshooting": ASSET_DIR / "troubleshooting.png",
        "test_loop": ASSET_DIR / "test_loop.png",
    }
    make_architecture(paths["architecture"])
    make_data_flow(paths["data_flow"])
    make_run_workflow(paths["run_workflow"])
    make_troubleshooting(paths["troubleshooting"])
    make_test_loop(paths["test_loop"])
    return paths


def east_asia(run, name="Microsoft YaHei"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)


def setup_doc(title, subtitle):
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(1.7)
    section.bottom_margin = Cm(1.6)
    section.left_margin = Cm(1.9)
    section.right_margin = Cm(1.9)
    normal = doc.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.line_spacing = 1.2
    normal.paragraph_format.space_after = Pt(5)
    for style_name, size in [("Heading 1", 15), ("Heading 2", 12), ("Heading 3", 10.5)]:
        style = doc.styles[style_name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(ACCENT)
        style.paragraph_format.space_before = Pt(9)
        style.paragraph_format.space_after = Pt(4)
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(title)
    east_asia(run)
    run.bold = True
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor.from_string(ACCENT)
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(subtitle)
    east_asia(run)
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(80, 91, 102)
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = info.add_run("项目：YVM-SLAM    文档版本：V3.0 图文版    日期：2026年4月30日")
    east_asia(run)
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(100, 100, 100)
    doc.add_page_break()
    return doc


def h(doc, text, level=1):
    doc.add_heading(text, level=level)


def p(doc, text):
    para = doc.add_paragraph()
    run = para.add_run(text)
    east_asia(run)
    return para


def bullets(doc, items):
    for item in items:
        para = doc.add_paragraph(style="List Bullet")
        para.paragraph_format.space_after = Pt(2)
        run = para.add_run(item)
        east_asia(run)


def numbered(doc, items):
    for item in items:
        para = doc.add_paragraph(style="List Number")
        para.paragraph_format.space_after = Pt(2)
        run = para.add_run(item)
        east_asia(run)


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell(cell, text, bold=False):
    cell.text = ""
    para = cell.paragraphs[0]
    run = para.add_run(str(text))
    east_asia(run)
    run.bold = bold
    run.font.size = Pt(9)


def small_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, header in enumerate(headers):
        set_cell(table.rows[0].cells[i], header, True)
        shade_cell(table.rows[0].cells[i], LIGHT)
    for row in rows:
        cells = table.add_row().cells
        for i, text in enumerate(row):
            set_cell(cells[i], text)
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    doc.add_paragraph()


def code(doc, text):
    para = doc.add_paragraph()
    p_pr = para._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), GRAY)
    p_pr.append(shd)
    run = para.add_run(text)
    run.font.name = "Consolas"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(9)


def picture(doc, path, caption):
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run()
    run.add_picture(str(path), width=Cm(16.2))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(caption)
    east_asia(r)
    r.italic = True
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(90, 90, 90)


def footer(doc):
    for section in doc.sections:
        para = section.footer.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run("YVM-SLAM 项目文档")
        east_asia(run)
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(120, 120, 120)


def build_design(diagrams):
    doc = setup_doc("YVM-SLAM 设计文档", "架构、数据流和关键模块说明")
    h(doc, "1. 项目定位")
    p(doc, "YVM-SLAM 是在 ORB-SLAM3 基础上扩展的动态环境 RGB-D SLAM。它的核心思路不是重新实现 SLAM，而是在 ORB-SLAM3 的特征提取入口处引入语义 mask，把人、车等动态区域中的关键点过滤掉，让后面的匹配、位姿估计、建图和回环尽量依赖静态背景。")
    p(doc, "当前代码还加入了 ROS2 数据桥、语义对象地图和 GTSAM 对象位置优化。这样项目既能离线跑 TUM RGB-D，也能接入外部分割节点，观察动态过滤对轨迹稳定性的影响。")
    picture(doc, diagrams["architecture"], "图 1  系统总体架构")

    h(doc, "2. 单帧处理流程")
    p(doc, "每一帧的处理可以理解为“先拿到图像和语义信息，再过滤动态特征，最后交给 ORB-SLAM3 主流程”。下面这张图比表格更直观：mask/detections 在 Tracking::GrabImageRGBD 中进入系统，真正改变特征的是 ORBextractor。")
    picture(doc, diagrams["data_flow"], "图 2  RGB-D 单帧处理流程")
    numbered(doc, [
        "rgbd_tum 读取 RGB、Depth 和 timestamp。如果启用 --ros，同时接收最新 /dynamic_mask 与 /dynamic_detections。",
        "System::SetExternalMask 和 SetExternalDetections 把外部语义结果写入 YoloDetection。",
        "Tracking::GrabImageRGBD 调用 YoloDetection::Detect，得到 mvDynamicMask、mvDynamicArea 和 mmDetectMap。",
        "ORBextractor 在 FAST/ORB 特征阶段删除动态区域关键点，再继续描述子计算。",
        "Tracking::Track 继续执行 ORB-SLAM3 的初始化、跟踪、重定位、关键帧判断和局部地图跟踪。",
        "如果检测到 chair、table、sofa 等静态类目标，UpdateMapObjectsFromDetections 会用深度中位数估计其 3D 世界坐标。",
    ])

    h(doc, "3. 关键模块说明")
    h(doc, "3.1 System", 2)
    p(doc, "System 是整个工程的装配入口。它负责加载 ORBvoc.txt、读取 YAML、创建 Atlas、KeyFrameDatabase、FrameDrawer、MapDrawer、Tracking、LocalMapping、LoopClosing、Viewer、YoloDetection 和 ObjectOptimizer。外部模块不应该直接操作 Tracking 内部状态，而应通过 System::SetExternalMask、System::SetExternalDetections 这类接口传入语义数据。")
    h(doc, "3.2 Tracking", 2)
    p(doc, "Tracking 是主线程中的核心逻辑。RGB-D 路径在 GrabImageRGBD 中完成：复制 RGB、调用检测模块、把动态 mask 写入 ORBextractor、把 RGB 转灰度、把深度转 CV_32F、构造 Frame、执行 Track、最后更新语义对象。这个函数是二次开发时最值得先读的入口。")
    h(doc, "3.3 YoloDetection 和 ORBextractor", 2)
    p(doc, "YoloDetection 当前更像一个“外部语义结果缓存器”，不是主要的 ONNX 推理入口。它把 ROS2 输入的 mask 转成 8UC1 二值图，把 detections 转成 label 到 bbox 的映射。ORBextractor 才是真正执行动态过滤的地方：只要关键点落在 mask=255 或动态 bbox 内，就从候选关键点列表中删除。")
    h(doc, "3.4 MapObject 和 ObjectOptimizer", 2)
    p(doc, "MapObject 保存语义对象的标签、ID、世界坐标、尺寸、观测次数和最近检测框。ObjectOptimizer 每 500ms 查看当前地图对象，如果某个对象积累了不少于 5 条观测，就用 GTSAM 对其位置做一次小规模优化。这个对象地图目前用于增强语义显示和对象级记录，不直接参与相机位姿优化。")

    h(doc, "4. 少量关键接口")
    small_table(doc, ["接口", "作用", "开发注意点"], [
        ("System::TrackRGBD", "处理一帧 RGB-D 并返回 Tcw", "Depth 必须与 RGB 对齐，时间戳要递增。"),
        ("System::SetExternalMask", "注入外部动态 mask", "推荐 mono8，非零像素表示动态区域。"),
        ("System::SetExternalDetections", "注入外部检测框", "bbox 坐标必须对应当前 RGB 尺寸。"),
        ("YoloDetection::Detect", "生成动态 mask/区域和检测结果", "调用后会清空缓存，避免旧消息长期生效。"),
        ("Tracking::UpdateMapObjectsFromDetections", "根据静态 bbox 更新对象地图", "只处理静态白名单标签。"),
    ])

    h(doc, "5. 当前设计中的风险点")
    bullets(doc, [
        "mask 访问建议增加边界判断。当前 ORBextractor 直接按关键点坐标访问 mask，若尺寸不一致或坐标异常，有越界风险。",
        "mask 和图像没有严格时间同步。ROS2 模式下代码使用“最新 mask”，极端情况下可能拿上一帧或下一帧的 mask。",
        "VLM 服务地址写死在代码中。实际部署时建议改为 YAML 参数或命令行参数。",
        "对象匹配只看标签和 1m 空间距离。同类物体密集时可能把两个对象合并。",
        "ROS2 依赖和 rgbd_tum 目标耦合较强。如果只想离线运行，建议后续拆分无 ROS 目标和 ROS2 目标。",
    ])
    footer(doc)
    path = OUT / "YVM-SLAM_设计文档.docx"
    doc.save(path)
    return path


def build_manual(diagrams):
    doc = setup_doc("YVM-SLAM 说明文档", "安装、运行、配置和排查指南")
    h(doc, "1. 先选运行模式")
    p(doc, "建议第一次使用时不要直接接 YOLO 或 ROS2。先跑离线基线，确认数据、词袋、YAML 和 ORB-SLAM3 主流程没问题；再启用 ROS2；最后接入外部分割节点。这样每一步出问题都能定位。")
    picture(doc, diagrams["run_workflow"], "图 1  推荐运行流程")

    h(doc, "2. 环境准备")
    bullets(doc, [
        "基础环境：Ubuntu 20.04 或兼容 Linux，GCC/G++，CMake。",
        "视觉依赖：OpenCV 4.x、Eigen3、Pangolin。",
        "SLAM 依赖：DBoW2、g2o、Sophus、Boost serialization、OpenSSL。",
        "扩展依赖：GTSAM、ONNX Runtime。当前 CMakeLists.txt 中 ONNXRUNTIME_DIR 是固定路径，换机器后要改。",
        "ROS2 模式依赖：ament_cmake、rclcpp、nav_msgs、sensor_msgs、vision_msgs、geometry_msgs、tf2、tf2_ros。",
        "数据：TUM RGB-D 序列、association 文件、对应 YAML。"
    ])

    h(doc, "3. 构建")
    p(doc, "推荐使用 build.sh，因为它会依次构建 DBoW2、g2o、Sophus，解压 ORB 词袋，然后构建 rgbd_tum。")
    code(doc, "cd YVM-SLAM\nsource /opt/ros/<distro>/setup.bash   # 如果使用 ROS2\nchmod +x build.sh\n./build.sh")
    p(doc, "如果 build.sh 失败，先不要急着改代码。先看失败位置：是第三方库失败、词袋解压失败，还是主项目链接失败。ONNX Runtime、GTSAM 和 ROS2 包路径是最常见问题。")
    small_table(doc, ["现象", "优先检查"], [
        ("找不到 onnxruntime.so", "CMakeLists.txt 里的 ONNXRUNTIME_DIR 是否指向真实目录。"),
        ("找不到 rclcpp/vision_msgs", "是否 source ROS2 setup.bash，是否安装对应消息包。"),
        ("Vocabulary loaded 前退出", "Vocabulary/ORBvoc.txt 是否存在。"),
        ("Pangolin 报错", "显示环境和 Pangolin 版本。"),
    ])

    h(doc, "4. 离线运行 TUM RGB-D")
    code(doc, "./Examples/RGB-D/rgbd_tum \\\n  Vocabulary/ORBvoc.txt \\\n  Examples/RGB-D/TUM1.yaml \\\n  /data/rgbd_dataset_freiburg1_xyz \\\n  /data/rgbd_dataset_freiburg1_xyz/associations.txt")
    p(doc, "运行成功时，一般会看到 ORB Vocabulary 加载成功、Start processing sequence、Viewer 窗口启动，结束后根目录生成 CameraTrajectory.txt 和 KeyFrameTrajectory.txt。")
    bullets(doc, [
        "如果提示 No images found，优先检查 association 文件里的相对路径是否能与数据集根目录拼起来。",
        "如果轨迹比例明显不对，优先检查 RGBD.DepthMapFactor。",
        "如果频繁 LOST，先确认内参和深度尺度，再考虑动态 mask。"
    ])

    h(doc, "5. ROS2 语义模式")
    p(doc, "带 --ros 运行后，rgbd_tum 会发布 RGB、Depth、CameraInfo、/odom 和 TF，同时订阅外部 /dynamic_mask 与 /dynamic_detections。外部分割节点可以订阅 /color/image，处理后再把 mask 发回来。")
    code(doc, "./Examples/RGB-D/rgbd_tum Vocabulary/ORBvoc.txt Examples/RGB-D/TUM1.yaml /data/tum_seq associations.txt --ros")
    bullets(doc, [
        "/dynamic_mask 推荐 mono8 或 8UC1。0 表示保留，非零会被二值化为动态区域。",
        "/dynamic_detections 使用 vision_msgs/Detection2DArray，class_id 会作为对象标签。",
        "bbox 坐标必须对应 RGB 图像坐标。如果外部分割节点缩放过图像，要把 bbox 映射回来。",
        "当前程序使用最新 mask，没有严格按 timestamp 同步，所以分割节点延迟过大时可能出现错位。"
    ])

    h(doc, "6. 参数怎么调")
    p(doc, "调参时先动最关键的几项，不要一次改很多。相机内参和 DepthMapFactor 是基础；ORB 参数影响速度和稳定性；mask 阈值与分割节点有关。")
    small_table(doc, ["参数", "什么时候调", "影响"], [
        ("Camera1.fx/fy/cx/cy", "换相机或换数据集", "直接影响位姿和对象位置。"),
        ("RGBD.DepthMapFactor", "轨迹尺度不对", "深度单位换算，TUM 常见 5000/5208。"),
        ("ORBextractor.nFeatures", "纹理少或速度慢", "调大更稳但更慢，调小更快但易丢。"),
        ("ORBextractor.iniThFAST", "低纹理场景", "调低能提更多点，但噪声会增加。"),
    ])

    h(doc, "7. 问题排查")
    picture(doc, diagrams["troubleshooting"], "图 2  常见问题排查路径")
    p(doc, "实际排查时建议按图中的顺序来：先判断是构建问题、运行问题还是效果问题。构建问题看依赖，运行问题看路径和配置，效果问题看内参、深度尺度、mask 对齐和时间同步。")

    h(doc, "8. 交付时建议保留的材料")
    bullets(doc, [
        "构建日志或成功截图。",
        "运行命令、数据集名称、YAML 文件。",
        "CameraTrajectory.txt 和 KeyFrameTrajectory.txt。",
        "Viewer 截图，最好能看到轨迹、地图点、mask 叠加或对象框。",
        "ROS2 topic list/topic hz 结果。",
        "ATE 评估结果，至少包含无 mask 和有 mask 两组对比。"
    ])
    footer(doc)
    path = OUT / "YVM-SLAM_说明文档.docx"
    doc.save(path)
    return path


def build_test(diagrams):
    doc = setup_doc("YVM-SLAM 测试文档", "构建、功能、集成和验收测试方案")
    h(doc, "1. 测试思路")
    p(doc, "测试不要只看“能不能跑起来”。这个项目至少要验证四件事：能构建、能跑离线 RGB-D、能接收动态 mask、能输出可评估的轨迹。语义对象地图和 GTSAM 优化属于增强功能，也要做最小验证。")
    picture(doc, diagrams["test_loop"], "图 1  测试闭环")

    h(doc, "2. 测试阶段")
    numbered(doc, [
        "构建验证：从空 build 目录重新 cmake 和 make，确认依赖不是靠缓存蒙混过关。",
        "离线基线：不带 --ros 跑短 TUM 序列，确认轨迹文件能生成。",
        "ROS2 发布：带 --ros 跑同一序列，检查 /color/image、/depth/image、/odom 和 TF。",
        "mask 集成：先用固定矩形 mask，再换真实分割节点，观察动态区域是否被过滤。",
        "detections 集成：发布 chair/table 等静态对象检测框，确认 MapObject 被创建和更新。",
        "轨迹评估：用 ground truth 评估 ATE，并记录无 mask/有 mask 对比。",
        "回归复测：每次改 ORBextractor、Tracking、CMake 或 ROS2 话题后，重复固定用例。"
    ])

    h(doc, "3. 关键测试用例")
    small_table(doc, ["编号", "用例", "通过标准"], [
        ("B-01", "构建 rgbd_tum", "生成 Examples/RGB-D/rgbd_tum，无链接错误。"),
        ("F-01", "离线 TUM 序列", "完整处理序列并生成 CameraTrajectory.txt。"),
        ("R-01", "ROS2 话题发布", "/color/image、/depth/image、/odom 有消息。"),
        ("M-01", "固定矩形 mask", "Viewer 可见 mask 叠加，动态区域特征减少。"),
        ("M-02", "真实分割 mask", "动态人体区域被过滤，程序不崩溃。"),
        ("O-01", "静态对象 bbox", "chair/table bbox 能新增或更新 MapObject。"),
        ("E-01", "错误路径/错误 YAML", "程序给出清楚错误，不出现无意义崩溃。"),
    ])

    h(doc, "4. 推荐测试命令")
    code(doc, "# 构建\nmkdir -p build && cd build\ncmake -DCMAKE_POLICY_VERSION_MINIMUM=3.5 -DCMAKE_BUILD_TYPE=Release ..\nmake rgbd_tum -j8\n\n# 离线基线\n./Examples/RGB-D/rgbd_tum Vocabulary/ORBvoc.txt Examples/RGB-D/TUM1.yaml /data/tum_seq /data/tum_seq/associations.txt\n\n# ROS2 模式\n./Examples/RGB-D/rgbd_tum Vocabulary/ORBvoc.txt Examples/RGB-D/TUM1.yaml /data/tum_seq /data/tum_seq/associations.txt --ros\n\n# 轨迹评估\npython evaluation/evaluate_ate_scale.py groundtruth.txt CameraTrajectory.txt")

    h(doc, "5. mask 测试要点")
    p(doc, "mask 测试最容易“看起来接上了，实际没生效”。建议先发布一个固定矩形 mask，这样能直观看到 Viewer 叠加区域，也能确认 ORB 特征是否减少。固定 mask 通过后，再接真实分割模型。")
    bullets(doc, [
        "mask 编码测试：mono8、8UC1、bgr8、rgb8 至少选两种验证。",
        "尺寸测试：同尺寸 mask 和缩放 mask 都要跑，确认 resize 后位置合理。",
        "极端测试：全 0 mask 应等同无过滤；全 255 mask 可能跟踪失败，但不应崩溃。",
        "延迟测试：降低分割节点频率，观察旧 mask 是否导致动态区域错位。"
    ])

    h(doc, "6. 结果记录")
    p(doc, "每次测试至少记录命令、数据集、配置文件、是否启用 --ros、mask 来源、平均跟踪时间、轨迹评估结果和问题现象。这样后续改代码时才能做真正的回归。")
    small_table(doc, ["记录项", "示例"], [
        ("数据集", "fr3/walking_xyz，前 1000 帧。"),
        ("配置", "Examples/RGB-D/TUM3.yaml，DepthMapFactor=5000。"),
        ("mask 来源", "固定矩形 / YOLO-World / YOLOv8-Seg。"),
        ("ATE RMSE", "无 mask: 0.xx，有 mask: 0.xx。"),
        ("平均跟踪时间", "mean/median。"),
        ("结论", "通过、失败或需要复测。"),
    ])

    h(doc, "7. 验收清单")
    bullets(doc, [
        "Release 构建成功，rgbd_tum 可执行文件存在。",
        "离线 TUM 序列可跑完，轨迹文件生成。",
        "ROS2 模式下关键话题正常发布。",
        "动态 mask 能进入系统并影响 ORB 特征过滤。",
        "静态检测框能生成 MapObject，对象观测数可增加。",
        "至少完成一次无 mask / 有 mask 的 ATE 对比。",
        "已记录失败用例和修复建议。"
    ])
    footer(doc)
    path = OUT / "YVM-SLAM_测试文档.docx"
    doc.save(path)
    return path


def main():
    diagrams = generate_diagrams()
    outputs = [build_design(diagrams), build_manual(diagrams), build_test(diagrams)]
    for path in outputs:
        print(path)


if __name__ == "__main__":
    main()
