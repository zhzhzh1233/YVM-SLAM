# -*- coding: utf-8 -*-
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs"
ASSETS = OUT / "assets"
OUT.mkdir(exist_ok=True)

ACCENT = "1F4E79"
LIGHT = "EAF2F8"
GRAY = "F5F7F9"


def east(run, name="Microsoft YaHei"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)


def setup(title, subtitle):
    doc = Document()
    s = doc.sections[0]
    s.top_margin = Cm(1.7)
    s.bottom_margin = Cm(1.6)
    s.left_margin = Cm(1.9)
    s.right_margin = Cm(1.9)

    normal = doc.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.line_spacing = 1.22
    normal.paragraph_format.space_after = Pt(5)

    for style_name, size in [("Heading 1", 15), ("Heading 2", 12), ("Heading 3", 10.5)]:
        st = doc.styles[style_name]
        st.font.name = "Microsoft YaHei"
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = RGBColor.from_string(ACCENT)
        st.paragraph_format.space_before = Pt(8)
        st.paragraph_format.space_after = Pt(4)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title)
    east(r)
    r.bold = True
    r.font.size = Pt(23)
    r.font.color.rgb = RGBColor.from_string(ACCENT)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(subtitle)
    east(r)
    r.font.size = Pt(12)
    r.font.color.rgb = RGBColor(80, 91, 102)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("YVM-SLAM 项目文档 | 精简实用版 | 2026年4月30日")
    east(r)
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(120, 120, 120)

    doc.add_page_break()
    return doc


def h(doc, text, level=1):
    doc.add_heading(text, level=level)


def p(doc, text):
    para = doc.add_paragraph()
    run = para.add_run(text)
    east(run)
    return para


def bullets(doc, items):
    for item in items:
        para = doc.add_paragraph(style="List Bullet")
        para.paragraph_format.space_after = Pt(2)
        run = para.add_run(item)
        east(run)


def numbered(doc, items):
    for item in items:
        para = doc.add_paragraph(style="List Number")
        para.paragraph_format.space_after = Pt(2)
        run = para.add_run(item)
        east(run)


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def table(doc, headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    for i, text in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = ""
        shade(cell, LIGHT)
        r = cell.paragraphs[0].add_run(text)
        east(r)
        r.bold = True
        r.font.size = Pt(9)
    for row in rows:
        cells = t.add_row().cells
        for i, text in enumerate(row):
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            cells[i].text = ""
            r = cells[i].paragraphs[0].add_run(str(text))
            east(r)
            r.font.size = Pt(9)
    doc.add_paragraph()


def code(doc, text):
    para = doc.add_paragraph()
    p_pr = para._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), GRAY)
    p_pr.append(shd)
    r = para.add_run(text)
    r.font.name = "Consolas"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    r.font.size = Pt(9)


def picture(doc, path, caption):
    if not path.exists():
        return
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.add_run().add_picture(str(path), width=Cm(15.6))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(caption)
    east(r)
    r.italic = True
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(95, 95, 95)


def footer(doc):
    for section in doc.sections:
        para = section.footer.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = para.add_run("YVM-SLAM 项目文档")
        east(r)
        r.font.size = Pt(8)
        r.font.color.rgb = RGBColor(120, 120, 120)


def build_design():
    doc = setup("YVM-SLAM 设计文档", "代码结构、核心流程和功能实现说明")

    h(doc, "1. 这个项目实现了什么")
    p(doc, "YVM-SLAM 是基于 ORB-SLAM3 改出来的 RGB-D 动态环境 SLAM。原来的 ORB-SLAM3 主要依赖 ORB 特征点，如果画面里有人、车等动态物体，这些动态点会干扰位姿估计。本项目的做法是在 ORB 特征提取阶段引入动态 mask 或检测框，把动态区域里的关键点删掉，再让后面的跟踪、建图和回环继续按 ORB-SLAM3 原流程执行。")
    p(doc, "简单说：它不是重写 SLAM，而是在 ORB-SLAM3 前端加了一层“动态区域过滤”，并额外增加了 ROS2 数据桥、语义对象记录和对象位置优化。")
    picture(doc, ASSETS / "data_flow.png", "图 1  单帧 RGB-D 处理流程")

    h(doc, "2. 代码主线")
    p(doc, "程序入口在 Examples/RGB-D/rgbd_tum.cc。它负责读 TUM RGB-D 数据集，创建 ORB_SLAM3::System，然后逐帧调用 SLAM.TrackRGBD(imRGB, imD, tframe)。如果命令行带 --ros，它还会发布 RGB、Depth、CameraInfo、/odom 和 TF，同时接收 /dynamic_mask 与 /dynamic_detections。")
    p(doc, "System 在 src/System.cc 中初始化整个系统：加载 ORB 词袋、读取配置、创建 Atlas、Tracking、LocalMapping、LoopClosing、Viewer、YoloDetection 和 ObjectOptimizer。代码主线可以简化为：rgbd_tum.cc 读帧 -> System::TrackRGBD -> Tracking::GrabImageRGBD -> ORBextractor 提取特征 -> Tracking::Track -> 输出位姿与轨迹。")
    table(doc, ["文件", "主要作用"], [
        ("Examples/RGB-D/rgbd_tum.cc", "程序入口；读数据集；可选启用 ROS2；调用 TrackRGBD。"),
        ("src/System.cc", "加载词袋和配置；创建各个模块和线程；提供 SetExternalMask 接口。"),
        ("src/Tracking.cc", "处理每一帧 RGB-D；调用动态过滤；执行 ORB-SLAM3 跟踪；更新对象地图。"),
        ("src/YoloDetect.cpp", "接收外部 mask/detections，整理成 ORBextractor 可用的数据。"),
        ("src/ORBextractor.cc", "提取 ORB 特征；删除动态 mask 或 bbox 内的关键点。"),
        ("src/MapObject.cc / ObjectOptimizer.cc", "保存语义对象，并用 GTSAM 优化对象位置。"),
    ])

    h(doc, "3. 动态区域过滤是怎么写的")
    p(doc, "动态过滤的核心在两个地方。第一处是 Tracking::GrabImageRGBD：它把当前 RGB 图传给 YoloDetection，调用 Detect 后，把 mpDetector->mvDynamicMask 和 mpDetector->mvDynamicArea 赋值给 mpORBextractorLeft。第二处是 ORBextractor::ComputeKeyPointsOctTree：ORBextractor 在得到 FAST 候选点后，会检查每个关键点是否落在动态 mask 或动态框里，如果是，就从候选关键点列表中 erase。")
    p(doc, "这样处理的好处是后续 Track、局部地图匹配、PnP、BA 等 ORB-SLAM3 原有逻辑不用大改，因为动态点在进入描述子和匹配之前就已经被过滤掉了。")
    p(doc, "YoloDetection::Detect 会把外部 mask 转成灰度二值图；没有 mask 时，则把检测框保存成动态区域。ORBextractor 删除关键点发生在描述子计算之前，因此动态区域的点不会参与后续匹配和位姿估计。")

    h(doc, "4. ROS2 功能是怎么接的")
    p(doc, "ROS2 逻辑主要写在 rgbd_tum.cc 里。程序带 --ros 运行时，会创建名为 orbslam_rgbd_publisher 的节点。它发布 /color/image、/depth/image、/color/camera_info、/depth/camera_info、/odom 和 TF；同时订阅 /dynamic_mask、/dynamic_detections、/yolo_world/active_prompts。")
    p(doc, "/dynamic_mask 会被 DecodeMaskImage 转成灰度 mask，然后通过 SLAM.SetExternalMask(mask) 写入检测模块。/dynamic_detections 会被转换成 map<string, vector<cv::Rect2i>>，再通过 SetExternalDetections 传进去。")
    p(doc, "位姿输出时，TrackRGBD 返回 Tcw，代码取 inverse 得到 Twc，然后发布到 /odom，并同步发布 odom 到 base_link 的 TF。")

    h(doc, "5. 语义对象功能")
    p(doc, "对象地图不是对所有检测目标都建图，而是只处理静态类，例如 chair、table、sofa、tv、laptop 等。Tracking::UpdateMapObjectsFromDetections 会取 bbox 中心区域的有效深度中位数，把 bbox 中心点反投影到相机坐标，再用当前位姿变换到世界坐标。")
    p(doc, "如果同类对象距离已有对象小于 1 米，就认为是同一个对象，追加一次观测；否则新建 MapObject。ObjectOptimizer 后台线程每 500ms 检查对象观测数，观测足够时用 GTSAM 对对象位置做优化。")
    h(doc, "6. 代码功能小结")
    bullets(doc, [
        "离线 RGB-D SLAM：读取 TUM RGB-D 数据，输出 CameraTrajectory.txt 和 KeyFrameTrajectory.txt。",
        "动态点过滤：接收 mask 或 bbox，在 ORBextractor 中删除动态区域关键点。",
        "ROS2 数据桥：发布 RGB、Depth、CameraInfo、/odom、TF，订阅外部语义结果。",
        "语义对象地图：根据静态检测框和深度估计对象 3D 位置。",
        "对象位置优化：后台用 GTSAM 对多帧对象观测做位置优化。",
    ])

    h(doc, "7. 目前需要注意的点")
    bullets(doc, [
        "ORBextractor 访问 mask 时最好补充边界判断，避免 mask 尺寸异常导致越界。",
        "ROS2 mask 当前使用“最新消息”，没有严格按时间戳同步，分割节点延迟过大时可能过滤错帧。",
        "ONNXRUNTIME_DIR 在 CMakeLists.txt 中是固定路径，换机器需要改。",
        "VLM 服务地址写在 Tracking::RequestVLM 中，实际部署建议做成配置项。",
        "对象匹配规则比较简单，同类物体很近时可能合并错误。"
    ])

    footer(doc)
    path = OUT / "YVM-SLAM_设计文档.docx"
    doc.save(path)
    return path


def build_manual():
    doc = setup("YVM-SLAM 说明文档", "如何构建、运行和排查问题")

    h(doc, "1. 推荐使用顺序")
    p(doc, "不要一开始就接 ROS2 和分割模型。建议先跑离线 TUM RGB-D，确认 ORB-SLAM3 主流程没问题；然后再带 --ros 跑，确认话题发布正常；最后接外部分割节点，测试 /dynamic_mask 是否真的影响特征过滤。")
    picture(doc, ASSETS / "run_workflow.png", "图 1  推荐运行流程")

    h(doc, "2. 构建")
    p(doc, "构建脚本 build.sh 会依次构建 DBoW2、g2o、Sophus，解压 ORBvoc.txt.tar.gz，然后构建 rgbd_tum。构建前需要确认 OpenCV、Eigen3、Pangolin、GTSAM、ONNX Runtime 和 ROS2 依赖都能被 CMake 找到。")
    code(doc, "cd YVM-SLAM\nsource /opt/ros/<distro>/setup.bash   # 使用 ROS2 时需要\nchmod +x build.sh\n./build.sh")
    p(doc, "如果只想手动构建主程序，可以进入 build 目录执行 cmake 和 make rgbd_tum。最常见的失败点是 ONNX Runtime 路径不对，或者 ROS2 环境没有 source。")
    code(doc, "mkdir -p build\ncd build\ncmake -DCMAKE_POLICY_VERSION_MINIMUM=3.5 -DCMAKE_BUILD_TYPE=Release ..\nmake rgbd_tum -j8")
    h(doc, "3. 离线运行")
    p(doc, "离线运行只需要词袋、YAML、TUM 序列目录和 association 文件。这个模式不依赖外部分割节点，适合先验证数据和相机参数。")
    code(doc, "./Examples/RGB-D/rgbd_tum \\\n  Vocabulary/ORBvoc.txt \\\n  Examples/RGB-D/TUM1.yaml \\\n  /data/rgbd_dataset_freiburg1_xyz \\\n  /data/rgbd_dataset_freiburg1_xyz/associations.txt")
    bullets(doc, [
        "看到 Vocabulary loaded，说明词袋路径正确。",
        "看到 Start processing sequence，说明 association 文件和图像路径基本正确。",
        "结束后生成 CameraTrajectory.txt 和 KeyFrameTrajectory.txt，说明轨迹保存正常。",
        "轨迹比例不对时，优先检查 RGBD.DepthMapFactor。"
    ])
    h(doc, "4. ROS2 和 mask 输入")
    p(doc, "带 --ros 运行后，程序会发布图像和里程计，也会订阅外部分割结果。外部分割节点一般订阅 /color/image，处理后发布 /dynamic_mask。YVM-SLAM 收到 mask 后，会在下一帧 Tracking 里把它传给 ORBextractor。")
    code(doc, "./Examples/RGB-D/rgbd_tum Vocabulary/ORBvoc.txt Examples/RGB-D/TUM1.yaml /data/tum_seq associations.txt --ros")
    table(doc, ["话题", "方向", "用途"], [
        ("/color/image", "发布", "给外部分割节点或 RViz 查看。"),
        ("/depth/image", "发布", "发布米制深度图。"),
        ("/odom", "发布", "输出当前相机位姿。"),
        ("/dynamic_mask", "订阅", "动态区域 mask，非零区域会过滤特征。"),
        ("/dynamic_detections", "订阅", "检测框和 label，用于对象地图。"),
    ])
    p(doc, "调试 mask 时可以先发布固定矩形 mask，确认 Viewer 上能看到叠加区域，再接真实分割节点。检测框主要用于对象地图，只有 chair、table 等静态标签会保存为 MapObject。")

    h(doc, "5. 关键参数")
    bullets(doc, [
        "Camera1.fx/fy/cx/cy：相机内参，换相机或换数据集必须改。",
        "RGBD.DepthMapFactor：深度缩放。TUM 常见 5000 或 5208；如果深度已经是米，就用 1。",
        "ORBextractor.nFeatures：每帧特征点数量。调大更稳但更慢，调小更快但可能丢跟踪。",
        "ORBextractor.iniThFAST/minThFAST：FAST 角点阈值。低纹理场景可以适当调低。",
        "Viewer 参数只影响显示，不影响轨迹估计。"
    ])
    h(doc, "6. 常见问题")
    bullets(doc, [
        "CMake 找不到 ONNX Runtime：修改 CMakeLists.txt 中的 ONNXRUNTIME_DIR。",
        "找不到 rclcpp 或 vision_msgs：先 source ROS2，再重新 cmake。",
        "No images found：检查 association 文件中的相对路径是否能和数据集根目录拼起来。",
        "mask 看起来没生效：先确认 /dynamic_mask 有消息、尺寸和 RGB 一致、动态区域像素非零。",
        "频繁 LOST：先无 mask 跑基线，确认内参和深度尺度正确，再检查 mask 是否错位。"
    ])

    footer(doc)
    path = OUT / "YVM-SLAM_说明文档.docx"
    doc.save(path)
    return path


def build_test():
    doc = setup("YVM-SLAM 测试文档", "测试步骤、检查点和验收标准")

    h(doc, "1. 测试目标")
    p(doc, "测试目标是确认项目能构建、能跑 TUM RGB-D、能接收动态 mask、能输出轨迹，并且对象地图的基本逻辑可用。测试时不要只看程序是否启动，还要检查轨迹文件、ROS2 话题、mask 是否生效、对象是否新增。")
    picture(doc, ASSETS / "test_loop.png", "图 1  测试闭环")

    h(doc, "2. 基础测试步骤")
    numbered(doc, [
        "从空 build 目录重新 cmake 和 make rgbd_tum，确认不是依赖缓存导致的假通过。",
        "不带 --ros 跑一段 TUM RGB-D 序列，确认能生成 CameraTrajectory.txt。",
        "带 --ros 跑同一序列，检查 /color/image、/depth/image、/odom 是否有消息。",
        "先发布一个固定矩形 /dynamic_mask，观察 Viewer 叠加和 ORB 特征数量变化。",
        "接入真实分割节点，比较无 mask 和有 mask 的轨迹表现。",
        "发布 chair/table 等静态检测框，确认 MapObject 能新增或更新。",
        "用 evaluation/evaluate_ate_scale.py 评估轨迹误差。"
    ])
    h(doc, "3. 重点用例")
    table(doc, ["用例", "检查点"], [
        ("构建测试", "rgbd_tum 可执行文件生成，无 ONNX/GTSAM/ROS2 链接错误。"),
        ("离线运行", "Viewer 正常启动，轨迹文件生成。"),
        ("ROS2 发布", "/color/image、/depth/image、/odom 有消息。"),
        ("mask 输入", "动态区域特征减少，程序不崩溃。"),
        ("detections 输入", "静态类检测框能更新 MapObject。"),
        ("异常路径", "词袋、YAML 或数据路径错误时有明确提示。"),
    ])

    h(doc, "4. 推荐命令")
    code(doc, "# 构建\nmkdir -p build && cd build\ncmake -DCMAKE_POLICY_VERSION_MINIMUM=3.5 -DCMAKE_BUILD_TYPE=Release ..\nmake rgbd_tum -j8\n\n# 离线运行\n./Examples/RGB-D/rgbd_tum Vocabulary/ORBvoc.txt Examples/RGB-D/TUM1.yaml /data/tum_seq /data/tum_seq/associations.txt\n\n# ROS2 模式\n./Examples/RGB-D/rgbd_tum Vocabulary/ORBvoc.txt Examples/RGB-D/TUM1.yaml /data/tum_seq /data/tum_seq/associations.txt --ros\n\n# 轨迹评估\npython evaluation/evaluate_ate_scale.py groundtruth.txt CameraTrajectory.txt")

    h(doc, "5. 验收标准")
    bullets(doc, [
        "Release 构建成功，Examples/RGB-D/rgbd_tum 存在。",
        "至少一个 TUM RGB-D 序列能跑完并生成轨迹。",
        "ROS2 模式下关键话题正常发布。",
        "/dynamic_mask 输入后，动态区域特征被过滤，程序不崩溃。",
        "静态对象检测框能生成或更新 MapObject。",
        "有轨迹评估结果，最好能和无 mask 基线对比。"
    ])

    footer(doc)
    path = OUT / "YVM-SLAM_测试文档.docx"
    doc.save(path)
    return path


def main():
    outputs = [build_design(), build_manual(), build_test()]
    for path in outputs:
        print(path)


if __name__ == "__main__":
    main()
