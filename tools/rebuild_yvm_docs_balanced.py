# -*- coding: utf-8 -*-
from pathlib import Path
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs"
ASSETS = OUT / "assets"
OUT.mkdir(exist_ok=True)

ACCENT = "1F4E79"
LIGHT = "EAF2F8"
GRAY = "F5F7F9"


def east(run, name="Microsoft YaHei"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)


def setup(title, subtitle):
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(1.7)
    section.bottom_margin = Cm(1.6)
    section.left_margin = Cm(1.9)
    section.right_margin = Cm(1.9)

    normal = doc.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.line_spacing = 1.2
    normal.paragraph_format.space_after = Pt(5)

    for style_name, size in [("Heading 1", 15), ("Heading 2", 12), ("Heading 3", 10.5)]:
        st = doc.styles[style_name]
        st.font.name = "Microsoft YaHei"
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = RGBColor.from_string(ACCENT)
        st.paragraph_format.space_before = Pt(8)
        st.paragraph_format.space_after = Pt(4)

    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(title)
    east(run)
    run.bold = True
    run.font.size = Pt(23)
    run.font.color.rgb = RGBColor.from_string(ACCENT)

    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(subtitle)
    east(run)
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(80, 91, 102)

    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run("YVM-SLAM 项目文档 | 适中详细版 | 2026年4月30日")
    east(run)
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(120, 120, 120)
    doc.add_page_break()
    return doc


def h(doc, text, level=1):
    doc.add_heading(text, level=level)


def p(doc, text):
    para = doc.add_paragraph()
    run = para.add_run(text)
    east(run)
    return para


def bullets(doc, items):
    for item in items:
        para = doc.add_paragraph(style="List Bullet")
        para.paragraph_format.space_after = Pt(2)
        run = para.add_run(item)
        east(run)


def numbered(doc, items):
    for item in items:
        para = doc.add_paragraph(style="List Number")
        para.paragraph_format.space_after = Pt(2)
        run = para.add_run(item)
        east(run)


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def table(doc, headers, rows):
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.style = "Table Grid"
    for i, text in enumerate(headers):
        cell = tbl.rows[0].cells[i]
        cell.text = ""
        shade(cell, LIGHT)
        run = cell.paragraphs[0].add_run(text)
        east(run)
        run.bold = True
        run.font.size = Pt(9)
    for row in rows:
        cells = tbl.add_row().cells
        for i, text in enumerate(row):
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            cells[i].text = ""
            run = cells[i].paragraphs[0].add_run(str(text))
            east(run)
            run.font.size = Pt(9)
    doc.add_paragraph()


def code(doc, text):
    para = doc.add_paragraph()
    p_pr = para._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), GRAY)
    p_pr.append(shd)
    run = para.add_run(text)
    run.font.name = "Consolas"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(9)


def picture(doc, name, caption):
    path = ASSETS / name
    if not path.exists():
        return
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.add_run().add_picture(str(path), width=Cm(15.6))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cap.add_run(caption)
    east(run)
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(95, 95, 95)


def footer(doc):
    for section in doc.sections:
        para = section.footer.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run("YVM-SLAM 项目文档")
        east(run)
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(120, 120, 120)


def build_design():
    doc = setup("YVM-SLAM 设计文档", "代码结构、核心流程和功能实现说明")

    h(doc, "1. 项目实现内容")
    p(doc, "YVM-SLAM 是在 ORB-SLAM3 基础上改造的 RGB-D 动态环境 SLAM。原始 ORB-SLAM3 使用 ORB 特征点进行跟踪和建图，当画面中出现人、车辆等动态目标时，这些动态目标上的特征点会干扰相机位姿估计。本项目的主要改动，就是在 ORB 特征提取阶段加入动态区域过滤。")
    p(doc, "系统接收外部语义分割结果，也就是 mask 或检测框。mask 表示图像中哪些区域属于动态目标；检测框是在没有像素级 mask 时的替代方案。进入 ORBextractor 后，动态区域内的关键点会被删除，剩下的静态背景特征再参与后续跟踪、建图和回环。")
    p(doc, "除了动态点过滤，项目还加入了 ROS2 数据桥、语义对象地图和对象位置优化。ROS2 部分负责发布图像、深度、里程计并订阅 mask；对象地图部分负责把静态物体检测框转换成世界坐标下的 MapObject。")
    picture(doc, "architecture.png", "图 1  系统总体架构")

    h(doc, "2. 代码调用主线")
    p(doc, "项目最主要的入口是 Examples/RGB-D/rgbd_tum.cc。它先解析命令行参数，再读取 association 文件，按顺序加载 RGB 图和深度图。每读到一帧，就调用 SLAM.TrackRGBD(imRGB, imD, tframe)。如果带 --ros 参数运行，它还会创建 ROS2 节点并发布/订阅相关话题。")
    p(doc, "System 在 src/System.cc 中负责初始化：加载 Vocabulary/ORBvoc.txt，读取 YAML 参数，创建 Atlas、KeyFrameDatabase、Tracking、LocalMapping、LoopClosing、Viewer、YoloDetection、ObjectOptimizer 等对象。可以把 System 理解成整个系统的装配器。")
    p(doc, "Tracking 是处理每一帧图像的核心。RGB-D 路径主要在 Tracking::GrabImageRGBD 中完成：先处理动态 mask，再把 RGB 转灰度，把深度转成 CV_32F，然后构造 Frame。Frame 构造过程中会调用 ORBextractor 提取特征。最后 Tracking::Track 执行 ORB-SLAM3 原有跟踪流程。")
    table(doc, ["模块/文件", "在代码中的作用"], [
        ("Examples/RGB-D/rgbd_tum.cc", "程序入口，读取数据集，处理 ROS2 话题，调用 TrackRGBD。"),
        ("src/System.cc", "初始化系统组件，创建线程，提供 SetExternalMask 等外部接口。"),
        ("src/Tracking.cc", "处理每帧 RGB-D，调用动态过滤，执行跟踪和对象地图更新。"),
        ("src/YoloDetect.cpp", "保存外部 mask/detections，并整理成动态区域数据。"),
        ("src/ORBextractor.cc", "提取 ORB 特征，并删除动态区域内的关键点。"),
        ("src/MapObject.cc", "保存语义对象的标签、位置、观测次数等信息。"),
    ])

    h(doc, "3. 动态过滤流程")
    p(doc, "动态过滤不是在后端优化阶段做的，而是在 ORB 特征提取阶段做的。Tracking::GrabImageRGBD 会调用 YoloDetection::Detect，得到 mvDynamicMask 或 mvDynamicArea，然后赋值给 mpORBextractorLeft。这样 ORBextractor 在提取关键点时，就能知道哪些区域应该跳过。")
    p(doc, "在 ORBextractor::ComputeKeyPointsOctTree 中，代码先用 FAST 得到候选关键点，再把关键点坐标映射到原图尺度。随后检查该点是否落在动态 mask 的白色区域，或者是否落在动态检测框内。如果落在动态区域，就从候选点列表中删除。删除完成后，剩余关键点再进入八叉树分配和描述子计算。")
    p(doc, "这种方式的优点是改动集中。ORB-SLAM3 后面的匹配、PnP、局部地图跟踪、BA 和回环检测基本不用改，因为动态点在进入这些环节前已经被过滤掉了。")
    picture(doc, "data_flow.png", "图 2  单帧 RGB-D 处理流程")

    h(doc, "4. ROS2 和对象地图")
    p(doc, "ROS2 代码集中在 rgbd_tum.cc。启用 --ros 后，程序发布 /color/image、/depth/image、/color/camera_info、/depth/camera_info、/odom 和 TF。同时订阅 /dynamic_mask、/dynamic_detections、/yolo_world/active_prompts。外部分割节点可以订阅 /color/image，处理后把 mask 发回 /dynamic_mask。")
    p(doc, "/dynamic_mask 会被 DecodeMaskImage 转成灰度 mask，然后通过 System::SetExternalMask 传入 YoloDetection。/dynamic_detections 会被转换成 label 到 bbox 的映射，再通过 SetExternalDetections 传入。位姿方面，TrackRGBD 返回 Tcw，程序取 inverse 得到 Twc，然后发布为 /odom 和 TF。")
    p(doc, "对象地图只处理静态类目标，例如 chair、table、sofa、tv、laptop。Tracking::UpdateMapObjectsFromDetections 会从检测框中心区域取有效深度中位数，将检测框中心反投影到三维相机坐标，再用当前相机位姿变换到世界坐标。如果同类对象距离已有对象小于 1 米，就更新已有对象；否则创建新的 MapObject。")
    p(doc, "ObjectOptimizer 是后台线程，每隔一段时间检查对象观测数量。如果一个对象已经有多次观测，就用 GTSAM 优化其三维位置。对象地图目前主要用于语义展示和对象级记录，没有直接参与相机位姿优化。")

    h(doc, "5. 功能与注意点")
    bullets(doc, [
        "离线 RGB-D SLAM：读取 TUM RGB-D 数据，输出 CameraTrajectory.txt 和 KeyFrameTrajectory.txt。",
        "动态点过滤：根据外部 mask 或检测框，在 ORBextractor 中删除动态区域关键点。",
        "ROS2 数据桥：发布 RGB、Depth、CameraInfo、/odom、TF，并接收外部分割结果。",
        "语义对象地图：把静态检测框和深度信息转换成世界坐标下的 MapObject。",
        "需要注意：ONNXRUNTIME_DIR 和 VLM 地址目前写在代码/配置中，换机器时要修改；mask 访问建议增加边界检查和时间同步。"
    ])

    h(doc, "6. 给初学者看的代码理解")
    p(doc, "如果不了解 SLAM，可以先把这个项目理解成一个“根据连续图像估计相机运动轨迹”的程序。RGB-D 表示每一帧不仅有彩色图，还有深度图。彩色图用来提取特征点，深度图用来把二维点转换成三维点。ORB-SLAM3 会不断匹配相邻帧和地图中的特征点，从而估计相机位姿。")
    p(doc, "YVM-SLAM 增加的部分，是在特征点进入 SLAM 之前做一次筛选。比如画面里有人走过，如果人的衣服上有很多 ORB 特征点，普通 ORB-SLAM3 可能把这些点也当成环境点。YVM-SLAM 通过外部分割 mask 知道“这个区域是动态目标”，于是把这个区域的关键点删掉，尽量只留下墙、地面、桌子等静态背景上的点。")
    p(doc, "从变量流向看，一帧数据先以 imRGB 和 imD 的形式进入 Tracking::GrabImageRGBD。imRGB 复制给 YoloDetection 做语义处理，同时会转成灰度图 mImGray。imD 会根据 RGBD.DepthMapFactor 转成 CV_32F 深度图。随后 Frame 构造函数接收灰度图、深度图、ORBextractor 和相机参数，生成当前帧的关键点、描述子和三维信息。")
    p(doc, "如果启用了 ROS2 mask，mask 不会直接改变图像本身，而是保存在 ORBextractor 的 mvDynamicMask 中。ORBextractor 提取关键点时会查这张 mask：如果关键点坐标对应的 mask 像素是动态区域，就删除这个关键点。这样后续描述子、匹配、位姿估计都看不到这个动态点。")

    h(doc, "7. 线程和模块之间的关系")
    p(doc, "System 初始化后会启动多个模块。Tracking 在主线程中随每帧图像运行；LocalMapping 在后台维护局部地图；LoopClosing 负责回环检测；Viewer 用 Pangolin 显示地图；ObjectOptimizer 负责语义对象位置优化。对于初学者来说，最重要的是先看 Tracking，因为每一帧的输入、动态过滤和位姿输出都从这里经过。")
    p(doc, "LocalMapping 和 LoopClosing 属于 ORB-SLAM3 原有后端，YVM-SLAM 没有大幅改写它们。也就是说，本项目的创新点主要在前端特征过滤和语义对象记录，而不是重新设计整个 SLAM 后端。理解这一点后，读代码会清楚很多：先看新增模块，再回到 ORB-SLAM3 原模块。")

    h(doc, "8. 阅读源码建议")
    numbered(doc, [
        "先读 Examples/RGB-D/rgbd_tum.cc，弄清楚命令行参数、图像读取、ROS2 话题和 TrackRGBD 调用。",
        "再读 src/System.cc 的构造函数，理解系统创建了哪些对象和线程。",
        "重点读 src/Tracking.cc 的 GrabImageRGBD，理解一帧 RGB-D 图像如何被处理。",
        "接着读 src/YoloDetect.cpp，看 mask 和 detections 如何转换成动态区域。",
        "最后读 src/ORBextractor.cc 中动态关键点删除逻辑，看动态过滤真正发生的位置。",
        "如果关心语义对象，再读 MapObject.cc、Map.cc 和 ObjectOptimizer.cc。"
    ])

    footer(doc)
    path = OUT / "YVM-SLAM_设计文档.docx"
    doc.save(path)
    return path


def build_manual():
    doc = setup("YVM-SLAM 说明文档", "构建、运行、参数和排查说明")

    h(doc, "1. 推荐运行顺序")
    p(doc, "建议按从简单到复杂的顺序运行。第一步先不启用 ROS2，只跑 TUM RGB-D 离线序列，确认词袋、相机参数和 ORB-SLAM3 主流程正常。第二步带 --ros 运行，检查图像、深度、/odom 和 TF 是否正常发布。第三步再接外部分割节点，验证 /dynamic_mask 是否真正影响 ORB 特征过滤。")
    picture(doc, "run_workflow.png", "图 1  推荐运行流程")

    h(doc, "2. 构建项目")
    p(doc, "构建脚本 build.sh 会构建 Thirdparty/DBoW2、Thirdparty/g2o、Thirdparty/Sophus，解压 ORB 词袋，然后在 build 目录构建 rgbd_tum。构建前需要准备 OpenCV、Eigen3、Pangolin、GTSAM、ONNX Runtime 和 ROS2 依赖。")
    code(doc, "cd YVM-SLAM\nsource /opt/ros/<distro>/setup.bash   # 使用 ROS2 时需要\nchmod +x build.sh\n./build.sh")
    p(doc, "如果手动构建，可以执行下面命令。构建成功后，应能在 Examples/RGB-D 下看到 rgbd_tum 可执行文件。如果找不到 ONNX Runtime，优先检查 CMakeLists.txt 中的 ONNXRUNTIME_DIR。")
    code(doc, "mkdir -p build\ncd build\ncmake -DCMAKE_POLICY_VERSION_MINIMUM=3.5 -DCMAKE_BUILD_TYPE=Release ..\nmake rgbd_tum -j8")

    h(doc, "3. 离线运行")
    p(doc, "离线模式需要四个参数：词袋路径、YAML 配置、TUM 序列目录、association 文件。这个模式不需要分割节点，适合先验证基础 SLAM 能力。")
    code(doc, "./Examples/RGB-D/rgbd_tum \\\n  Vocabulary/ORBvoc.txt \\\n  Examples/RGB-D/TUM1.yaml \\\n  /data/rgbd_dataset_freiburg1_xyz \\\n  /data/rgbd_dataset_freiburg1_xyz/associations.txt")
    bullets(doc, [
        "看到 Vocabulary loaded，说明词袋加载成功。",
        "看到 Start processing sequence，说明图像路径和 association 文件基本正确。",
        "程序结束后生成 CameraTrajectory.txt 和 KeyFrameTrajectory.txt，说明轨迹保存正常。",
        "如果轨迹比例不对，优先检查 RGBD.DepthMapFactor。"
    ])

    h(doc, "4. ROS2 和 mask 输入")
    p(doc, "带 --ros 运行后，程序会发布 RGB 图像、深度图、相机信息和里程计，同时订阅动态 mask 和检测框。外部分割节点通常订阅 /color/image，输出 /dynamic_mask。YVM-SLAM 收到 mask 后，会在下一帧 Tracking 中传给 ORBextractor。")
    code(doc, "./Examples/RGB-D/rgbd_tum Vocabulary/ORBvoc.txt Examples/RGB-D/TUM1.yaml /data/tum_seq associations.txt --ros")
    table(doc, ["话题", "方向", "用途"], [
        ("/color/image", "发布", "给外部分割节点或 RViz 查看。"),
        ("/depth/image", "发布", "发布米制深度图。"),
        ("/odom", "发布", "输出当前相机位姿。"),
        ("/dynamic_mask", "订阅", "动态区域 mask，非零区域会过滤特征。"),
        ("/dynamic_detections", "订阅", "检测框和 label，用于对象地图。"),
    ])
    p(doc, "调试 mask 时，建议先发布固定矩形 mask。如果 Viewer 能看到 mask 叠加，并且动态区域特征减少，说明 SLAM 这一侧逻辑正常。之后再接真实 YOLO 或 YOLO-World 分割节点。")

    h(doc, "5. 参数和结果查看")
    bullets(doc, [
        "Camera1.fx/fy/cx/cy：相机内参，换相机或换数据集必须修改。",
        "RGBD.DepthMapFactor：深度缩放因子，TUM 常见为 5000 或 5208。",
        "ORBextractor.nFeatures：每帧提取的 ORB 特征数，调大更稳但更慢。",
        "CameraTrajectory.txt：普通帧轨迹，用于 TUM ATE 评估。",
        "KeyFrameTrajectory.txt：关键帧轨迹，数量比普通帧少。",
        "Viewer：查看地图点、相机轨迹、mask 叠加和检测框显示是否正常。"
    ])

    h(doc, "6. 常见问题")
    p(doc, "如果构建失败，先看依赖是否能被 CMake 找到；如果运行失败，先看词袋、YAML 和数据路径；如果效果不好，先看相机内参、深度尺度和 mask 是否对齐。下面的图给出常见排查路径。")
    picture(doc, "troubleshooting.png", "图 2  常见问题排查路径")
    bullets(doc, [
        "找不到 rclcpp 或 vision_msgs：先 source ROS2，再重新 cmake。",
        "No images found：检查 association 文件中的路径是否能和数据集根目录拼起来。",
        "mask 没效果：检查 /dynamic_mask 是否有消息、尺寸是否和 RGB 一致、非零区域是否正确。",
        "频繁 LOST：先无 mask 跑基线，确认内参和深度尺度正确，再检查 mask 是否错位。"
    ])

    h(doc, "7. 项目目录怎么理解")
    p(doc, "src 和 include 是核心代码目录。src 里放实现文件，include 里放头文件。ORB-SLAM3 原有模块大多保留，例如 System、Tracking、LocalMapping、LoopClosing、ORBextractor、Map、KeyFrame、MapPoint 等。YVM-SLAM 新增或重点改动的文件主要包括 YoloDetect、MapObject、ObjectOptimizer，以及 Tracking 和 ORBextractor 中与动态过滤相关的部分。")
    p(doc, "Examples 目录是可执行示例。当前主要使用 Examples/RGB-D/rgbd_tum.cc，它既能跑 TUM RGB-D 离线数据，也能通过 --ros 打开 ROS2 发布和订阅。Vocabulary 目录放 ORB 词袋；Thirdparty 放 DBoW2、g2o、Sophus、ONNX Runtime 等依赖；models 目录放 YOLO 或 ONNX 模型文件；evaluation 目录放轨迹评估脚本。")

    h(doc, "8. 第一次运行时怎么看输出")
    p(doc, "程序启动后会先打印输入传感器类型，然后加载 ORB 词袋。如果卡在词袋加载，通常是 ORBvoc.txt 不存在或路径写错。词袋加载成功后，程序会读取数据集图片数量，并进入 Start processing sequence。这个阶段如果报 No images found，一般是 association 文件或序列根目录不匹配。")
    p(doc, "运行过程中 Viewer 会显示相机轨迹和地图点。正常情况下，静态场景里的地图点会逐渐稳定；如果地图点很少或轨迹很快丢失，先检查相机内参、图像清晰度、DepthMapFactor 和 ORB 特征数量。不要一开始就怀疑 mask，因为无 mask 基线没跑稳时，动态过滤也无法解决根本问题。")
    p(doc, "程序结束后会保存 CameraTrajectory.txt 和 KeyFrameTrajectory.txt。CameraTrajectory.txt 是普通帧轨迹，适合做 TUM ATE 评估；KeyFrameTrajectory.txt 是关键帧轨迹，数量更少，用来观察关键帧级别的运动。")

    h(doc, "9. mask 和 detections 的输入要求")
    p(doc, "mask 最好是 mono8 或 8UC1 编码，尺寸最好和 /color/image 一致。mask 中 0 表示保留区域，非零表示动态区域。YoloDetection 会把 mask 二值化，所以外部分割节点只要保证动态区域像素非零即可。")
    p(doc, "detections 使用 vision_msgs::msg::Detection2DArray。代码会读取 detection 的 class_id 和 bbox，把它转成 label 到 cv::Rect2i 的映射。动态区域过滤可以使用 bbox，但对象地图只处理静态标签，例如 chair、table、sofa、tv、laptop。如果发送 person，通常用于动态过滤，不会作为长期静态对象保存。")

    h(doc, "10. 常用调试顺序")
    numbered(doc, [
        "先不带 --ros 跑离线序列，确认能生成轨迹。",
        "再带 --ros 运行，用 ros2 topic list 和 ros2 topic hz 检查图像和 /odom。",
        "发布一个固定矩形 mask，确认 Viewer 中能看到 mask 叠加。",
        "接真实分割节点，确认 mask 尺寸、编码和延迟。",
        "最后再测试 detections 和对象地图。"
    ])

    footer(doc)
    path = OUT / "YVM-SLAM_说明文档.docx"
    doc.save(path)
    return path


def build_test():
    doc = setup("YVM-SLAM 测试文档", "测试步骤、检查点和验收标准")

    h(doc, "1. 测试目标")
    p(doc, "测试目标是确认项目能正常构建、能跑通 TUM RGB-D、能接收动态 mask、能输出轨迹，并且对象地图基本可用。测试时不要只看程序是否启动，还要检查轨迹文件、ROS2 话题、mask 是否生效、对象是否新增。")
    p(doc, "测试顺序建议从简单到复杂：先无 mask 跑基线，再固定 mask，再真实 mask。这样如果轨迹变差，可以判断问题是出在 SLAM 主流程、mask 接入，还是外部分割模型。")
    picture(doc, "test_loop.png", "图 1  测试闭环")

    h(doc, "2. 基础测试步骤")
    numbered(doc, [
        "从空 build 目录重新 cmake 和 make rgbd_tum，确认构建可重复。",
        "不带 --ros 跑一段 TUM RGB-D 序列，确认能生成 CameraTrajectory.txt。",
        "带 --ros 跑同一序列，检查 /color/image、/depth/image、/odom 是否有消息。",
        "发布固定矩形 /dynamic_mask，观察 Viewer 叠加和特征过滤效果。",
        "接入真实分割节点，比较无 mask 和有 mask 的轨迹表现。",
        "发布 chair/table 等静态检测框，确认 MapObject 能新增或更新。",
        "用 evaluation/evaluate_ate_scale.py 评估轨迹误差。"
    ])

    h(doc, "3. 重点用例")
    table(doc, ["用例", "检查点"], [
        ("构建测试", "rgbd_tum 可执行文件生成，无 ONNX/GTSAM/ROS2 链接错误。"),
        ("离线运行", "Viewer 正常启动，轨迹文件生成。"),
        ("ROS2 发布", "/color/image、/depth/image、/odom 有消息。"),
        ("mask 输入", "动态区域特征减少，程序不崩溃。"),
        ("detections 输入", "静态类检测框能更新 MapObject。"),
        ("异常路径", "词袋、YAML 或数据路径错误时有明确提示。"),
    ])

    h(doc, "4. mask 测试说明")
    p(doc, "mask 是本项目最关键的新增输入，需要单独测试。全 0 mask 应该等价于无过滤；固定矩形 mask 应该能在 Viewer 上看到叠加区域；全 255 mask 会删除大量特征，可能导致跟踪失败，但程序不应该崩溃。")
    p(doc, "真实分割节点测试时，要注意 mask 的尺寸、编码和延迟。尺寸不一致会导致过滤位置错；编码不对会导致 DecodeMaskImage 失败；延迟太大可能让当前帧使用旧 mask。")

    h(doc, "5. 推荐命令")
    code(doc, "# 构建\nmkdir -p build && cd build\ncmake -DCMAKE_POLICY_VERSION_MINIMUM=3.5 -DCMAKE_BUILD_TYPE=Release ..\nmake rgbd_tum -j8\n\n# 离线运行\n./Examples/RGB-D/rgbd_tum Vocabulary/ORBvoc.txt Examples/RGB-D/TUM1.yaml /data/tum_seq /data/tum_seq/associations.txt\n\n# ROS2 模式\n./Examples/RGB-D/rgbd_tum Vocabulary/ORBvoc.txt Examples/RGB-D/TUM1.yaml /data/tum_seq /data/tum_seq/associations.txt --ros\n\n# 轨迹评估\npython evaluation/evaluate_ate_scale.py groundtruth.txt CameraTrajectory.txt")

    h(doc, "6. 验收标准")
    bullets(doc, [
        "Release 构建成功，Examples/RGB-D/rgbd_tum 存在。",
        "至少一个 TUM RGB-D 序列能跑完并生成轨迹。",
        "ROS2 模式下关键话题正常发布。",
        "/dynamic_mask 输入后，动态区域特征被过滤，程序不崩溃。",
        "静态对象检测框能生成或更新 MapObject。",
        "有轨迹评估结果，最好能和无 mask 基线对比。"
    ])

    h(doc, "7. 测试记录建议")
    p(doc, "每次测试建议记录代码版本、数据集名称、YAML 文件、是否启用 --ros、mask 来源、运行命令、是否生成轨迹、ATE 结果和异常日志。这样后续修改代码后，可以用同一组条件做回归。")

    h(doc, "8. 具体功能应该怎么判断")
    p(doc, "构建功能是否正常，不能只看 cmake 是否结束，还要确认 Examples/RGB-D/rgbd_tum 是否生成。第一次构建还要确认 Vocabulary/ORBvoc.txt 已经解压，否则运行时会在加载词袋阶段失败。")
    p(doc, "离线 SLAM 是否正常，主要看三点：程序能处理完整序列，Viewer 中轨迹不是马上丢失，结束后生成 CameraTrajectory.txt。如果这三点不满足，先不要接 mask，应该先检查数据路径、YAML 和深度尺度。")
    p(doc, "ROS2 功能是否正常，主要看 topic。启动 --ros 后应该能看到 /color/image、/depth/image、/odom 等话题。/odom 中 position 和 orientation 应连续变化，不应该出现 NaN 或大幅跳变。")
    p(doc, "mask 功能是否正常，最直接的办法是先用固定矩形 mask 测试。如果固定 mask 都不能显示或不能影响特征数，说明问题在 YVM-SLAM 接收或处理 mask 的逻辑；如果固定 mask 正常但真实分割不正常，再去查外部分割节点。")

    h(doc, "9. 轨迹评估说明")
    p(doc, "如果数据集有 ground truth，可以使用 evaluation/evaluate_ate_scale.py 评估 CameraTrajectory.txt。建议至少跑两组：无 mask 基线和有 mask 版本。动态场景下，有 mask 的结果应该减少动态物体带来的干扰，但如果 mask 错位或延迟太大，也可能让结果变差。")
    p(doc, "评估时不要只记录一个误差数值，还要记录数据集、YAML、DepthMapFactor、mask 来源和运行命令。这样后续改代码时，才能复现同一组实验。")

    h(doc, "10. 回归测试建议")
    bullets(doc, [
        "修改 ORBextractor 后，必须重新测试固定 mask、全 0 mask 和全 255 mask。",
        "修改 Tracking::GrabImageRGBD 后，必须测试离线运行和 ROS2 mask 输入。",
        "修改 CMakeLists.txt 后，最好删除 build 目录重新构建。",
        "修改对象地图逻辑后，测试 chair/table 等静态检测框是否还能生成 MapObject。",
        "修改 ROS2 话题名后，同步检查外部分割节点和文档中的话题说明。"
    ])

    footer(doc)
    path = OUT / "YVM-SLAM_测试文档.docx"
    doc.save(path)
    return path


def main():
    for path in [build_design(), build_manual(), build_test()]:
        print(path)


if __name__ == "__main__":
    main()
